import re
from pathlib import Path

from docx import Document

try:
    from pypdf import PdfReader  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover
    PdfReader = None


class ResumeParser:
    def __init__(self, chunk_size=200, overlap=50):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def _clean_text(self, text):
        text = text.replace("\t", " ")
        text = re.sub(r" +", " ", text)
        text = re.sub(r"\n+", "\n", text)
        return text.strip()

    def extract_text(self, file_path):
        suffix = Path(file_path).suffix.lower()
        if suffix == ".docx":
            return self.extract_from_docx(file_path)
        if suffix == ".pdf":
            return self.extract_from_pdf(file_path)
        raise ValueError("仅支持 .docx 或 .pdf 格式")

    def extract_from_docx(self, file_path):
        doc = Document(file_path)
        content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    content.append(row_text)

        raw_text = "\n".join(content)
        return self._clean_text(raw_text)

    def extract_from_pdf(self, file_path):
        if PdfReader is None:
            raise ValueError("当前环境缺少 pypdf，无法解析 PDF 简历")

        reader = PdfReader(file_path)
        content = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                content.append(text)

        return self._clean_text("\n".join(content))

    def get_chunks(self, text):
        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = start + self.chunk_size
            chunks.append(text[start:end])
            start += self.chunk_size - self.overlap

            if start >= len(text):
                break

        return chunks
